"""
Document Ingestion Agent
Handles PDF, Excel, Word, Images — routes to appropriate parser
"""
import os
import io
import uuid
import logging
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from datetime import datetime

import pypdf
import openpyxl
import docx
from PIL import Image

from core.vectorstore import add_documents
from core.llm import get_embeddings

logger = logging.getLogger(__name__)

UPLOAD_DIR = os.getenv("UPLOAD_DIR", "./uploads")
CHUNK_SIZE = 800
CHUNK_OVERLAP = 100


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks"""
    if not text.strip():
        return []
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk)
        if end >= len(words):
            break
        start = end - overlap
    return chunks


def extract_pdf(file_path: str) -> Tuple[str, int, List[Dict]]:
    """Extract text from PDF — returns (full_text, page_count, pages)"""
    pages = []
    full_text = ""
    try:
        reader = pypdf.PdfReader(file_path)
        page_count = len(reader.pages)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({"page": i + 1, "text": text})
            full_text += f"\n--- Page {i + 1} ---\n{text}"
        return full_text, page_count, pages
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return "", 0, []


def extract_excel(file_path: str) -> Tuple[str, int, List[Dict]]:
    """Extract text from Excel"""
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        all_text = ""
        pages = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(c) if c is not None else "" for c in row]
                if any(v.strip() for v in row_vals):
                    rows.append("\t".join(row_vals))
            sheet_text = f"\n=== Sheet: {sheet_name} ===\n" + "\n".join(rows)
            all_text += sheet_text
            pages.append({"page": sheet_name, "text": sheet_text})
        return all_text, len(wb.sheetnames), pages
    except Exception as e:
        logger.error(f"Excel extraction error: {e}")
        return "", 0, []


def extract_docx(file_path: str) -> Tuple[str, int, List[Dict]]:
    """Extract text from Word document"""
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells)
                full_text += f"\n{row_text}"
        return full_text, 1, [{"page": 1, "text": full_text}]
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return "", 0, []


def extract_image(file_path: str) -> Tuple[str, int, List[Dict]]:
    """Extract text from image using OCR agent"""
    try:
        from agents.ocr_agent import ocr_image
        text = ocr_image(file_path)
        return text, 1, [{"page": 1, "text": text}]
    except Exception as e:
        logger.error(f"Image extraction error: {e}")
        return "", 0, []


def extract_csv(file_path: str) -> Tuple[str, int, List[Dict]]:
    """Extract text from CSV"""
    import pandas as pd
    try:
        df = pd.read_csv(file_path)
        text = df.to_string()
        return text, 1, [{"page": 1, "text": text}]
    except Exception as e:
        logger.error(f"CSV extraction error: {e}")
        return "", 0, []


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".xlsx": extract_excel,
    ".xls": extract_excel,
    ".docx": extract_docx,
    ".doc": extract_docx,
    ".png": extract_image,
    ".jpg": extract_image,
    ".jpeg": extract_image,
    ".csv": extract_csv,
}


async def ingest_document(
    file_path: str,
    document_id: int,
    original_name: str,
) -> Dict[str, Any]:
    """
    Main ingestion pipeline:
    1. Extract text based on file type
    2. Chunk text
    3. Generate embeddings
    4. Store in ChromaDB
    """
    ext = Path(file_path).suffix.lower()
    extractor = EXTRACTORS.get(ext)

    if not extractor:
        return {"success": False, "error": f"Unsupported file type: {ext}"}

    # Extract text
    logger.info(f"Extracting text from {original_name} ({ext})")
    full_text, page_count, pages = extractor(file_path)

    if not full_text.strip():
        # Try OCR for PDFs with no text
        if ext == ".pdf":
            from agents.ocr_agent import ocr_pdf
            full_text = ocr_pdf(file_path)
            pages = [{"page": 1, "text": full_text}]
            page_count = 1

    if not full_text.strip():
        return {"success": False, "error": "No text could be extracted"}

    # Chunk by page
    chunks = []
    chunk_metadata = []

    for page_info in pages:
        page_chunks = chunk_text(str(page_info["text"]))
        for j, chunk in enumerate(page_chunks):
            chunks.append(chunk)
            chunk_metadata.append({
                "document_id": str(document_id),
                "document_name": original_name,
                "page": str(page_info.get("page", 1)),
                "chunk_index": str(len(chunks)),
                "file_type": ext.lstrip("."),
            })

    if not chunks:
        return {"success": False, "error": "No chunks generated"}

    # Generate unique IDs
    chunk_ids = [f"doc_{document_id}_chunk_{i}" for i in range(len(chunks))]

    # Add to vector store
    add_documents(texts=chunks, ids=chunk_ids, metadatas=chunk_metadata)

    logger.info(f"✅ Ingested {original_name}: {page_count} pages, {len(chunks)} chunks")

    return {
        "success": True,
        "page_count": page_count,
        "chunk_count": len(chunks),
        "full_text": full_text[:2000],  # First 2000 chars for summary
    }
